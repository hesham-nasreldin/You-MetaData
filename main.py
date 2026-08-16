import streamlit as st
import yt_dlp
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO

strings = {
    "en": {
        "title": "YouTube Metadata",
        "urlLabel": "Video URL",
        "fetchBtn": "Get Metadata",
        "resultTitle": "Video Information",
        "labelTitle": "Title",
        "labelChannel": "Channel",
        "labelDate": "Upload Date",
        "labelLength": "Duration",
        "errorNoUrl": "Please provide a video URL.",
        "errorFetch": "Error fetching metadata.",
    },
    "ar": {
        "title": "بيانات الفيديو",
        "urlLabel": "رابط الفيديو",
        "fetchBtn": "جلب المعلومات",
        "resultTitle": "معلومات الفيديو",
        "labelTitle": "العنوان",
        "labelChannel": "القناة",
        "labelDate": "تاريخ الرفع",
        "labelLength": "المدة",
        "errorNoUrl": "يرجى إدخال رابط الفيديو.",
        "errorFetch": "حدث خطأ أثناء جلب البيانات.",
    },
}


def create_word_document(results, lang_code):
    """Create a Word document from the results."""
    doc = Document()

    # Add title
    title = doc.add_heading(strings[lang_code]["resultTitle"], 0)

    # Add table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "URL"
    header_cells[1].text = strings[lang_code]["labelTitle"]
    header_cells[2].text = strings[lang_code]["labelChannel"]
    header_cells[3].text = strings[lang_code]["labelDate"]
    header_cells[4].text = strings[lang_code]["labelLength"]

    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = result["url"]
        row_cells[1].text = result.get(strings[lang_code]["labelTitle"], "")
        row_cells[2].text = result.get(strings[lang_code]["labelChannel"], "")
        row_cells[3].text = result.get(strings[lang_code]["labelDate"], "")
        row_cells[4].text = result.get(strings[lang_code]["labelLength"], "")

        # Add error message if present
        if result.get("error"):
            error_para = row_cells[1].paragraphs[0]
            error_para.text = f"Error: {result['error']}"

    # Convert to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


st.set_page_config(page_title="YouTube Metadata", layout="centered")

# Language selector
lang_label = "Language / اللغة"
lang = st.selectbox(lang_label, options=["English", "العربية"])
lang_code = "ar" if lang == "العربية" else "en"

# Inject RTL and font for Arabic
if lang_code == "ar":
    st.markdown(
        """
        <link href="https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;700&display=swap" rel="stylesheet">
        <style>
        html, body, [data-testid="stAppViewContainer"] {direction: rtl; text-align: right; font-family: 'Cairo', sans-serif}
        .stButton>button {float: none}
        </style>
        """,
        unsafe_allow_html=True,
    )


st.title(strings[lang_code]["title"])

with st.form(key="meta_form"):
    url_text = st.text_area(
        label=strings[lang_code]["urlLabel"],
        value="",
        help="Enter one video URL per line",
        placeholder="https://www.youtube.com/watch?v=...\nhttps://youtu.be/...",
        height=120,
    )
    submit = st.form_submit_button(strings[lang_code]["fetchBtn"])

if submit:
    urls = [u.strip() for u in (url_text or "").splitlines() if u.strip()]
    if not urls:
        st.error(strings[lang_code]["errorNoUrl"])
    else:
        results = []
        ydl_opts = {"quiet": True, "extract_flat": True}
        for u in urls:
            with st.spinner(f"Fetching: {u}"):
                try:
                    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                        info = ydl.extract_info(u, download=False)

                    title = info.get("title") or "-"
                    channel = info.get("uploader") or "-"
                    raw_date = info.get("upload_date")
                    if raw_date and len(raw_date) == 8:
                        upload_date = f"{raw_date[:4]}-{raw_date[4:6]}-{raw_date[6:]}"
                    else:
                        upload_date = "-"

                    # Format duration
                    duration_seconds = info.get("duration")
                    if duration_seconds:
                        hours = duration_seconds // 3600
                        minutes = (duration_seconds % 3600) // 60
                        seconds = duration_seconds % 60
                        if hours > 0:
                            duration_str = f"{hours}:{minutes:02d}:{seconds:02d}"
                        else:
                            duration_str = f"{minutes}:{seconds:02d}"
                    else:
                        duration_str = "-"

                    results.append(
                        {
                            "url": u,
                            strings[lang_code]["labelTitle"]: title,
                            strings[lang_code]["labelChannel"]: channel,
                            strings[lang_code]["labelDate"]: upload_date,
                            strings[lang_code]["labelLength"]: duration_str,
                            "error": "",
                        }
                    )
                except yt_dlp.utils.DownloadError as e:
                    results.append(
                        {
                            "url": u,
                            strings[lang_code]["labelTitle"]: "",
                            strings[lang_code]["labelChannel"]: "",
                            strings[lang_code]["labelDate"]: "",
                            strings[lang_code]["labelLength"]: "",
                            "error": f"DownloadError: {e}",
                        }
                    )
                except Exception as e:
                    results.append(
                        {
                            "url": u,
                            strings[lang_code]["labelTitle"]: "",
                            strings[lang_code]["labelChannel"]: "",
                            strings[lang_code]["labelDate"]: "",
                            strings[lang_code]["labelLength"]: "",
                            "error": str(e),
                        }
                    )

        st.subheader(strings[lang_code]["resultTitle"])
        # Display results as a table
        st.dataframe(results)

        # Download button
        doc_bytes = create_word_document(results, lang_code)
        st.download_button(
            label=(
                "📥 Download as Word (.docx)"
                if lang_code == "en"
                else "📥 تحميل كملف Word"
            ),
            data=doc_bytes,
            file_name="youtube_metadata.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
